"""
03_build_report_humanized.py
Produces: Data_Analysis.docx
Format: A4, Arial 12, 1.5 line spacing, justified indented paragraphs.
Matches sample report style — paragraph-form narrative, italic Finding/Insight.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

BASE   = Path(__file__).parent
FIGS   = BASE / "figures"
OUTPUT = BASE / "Data_Analysis.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_margins(section, top=1.0, right=1.0, bottom=1.0, left=1.5):
    section.top_margin    = Inches(top)
    section.right_margin  = Inches(right)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)


def apply_font(run, bold=False, italic=False, size=12, color=None):
    run.font.name   = "Arial"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_format(para, spacing=1.5, space_before=0, space_after=6,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=None):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = spacing
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.alignment         = align
    if first_indent is not None:
        pf.first_line_indent = first_indent


def add_heading(doc, text):
    """Bold section or sub-section heading, left-aligned."""
    p = doc.add_paragraph()
    set_para_format(p, space_before=12, space_after=6,
                    align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    apply_font(run, bold=True)
    return p


def add_body(doc, text):
    """Justified paragraph with first-line indent — standard narrative text."""
    p = doc.add_paragraph()
    set_para_format(p, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_indent=Inches(0.5))
    run = p.add_run(text)
    apply_font(run)
    return p


def add_label_line(doc, text):
    """Plain left-aligned line for dataset field labels (Source:, Date:, etc.)."""
    p = doc.add_paragraph()
    set_para_format(p, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    apply_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_format(p, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    apply_font(run)
    return p


def add_numbered(doc, num, text):
    """Numbered list item matching sample style (1. text)."""
    p = doc.add_paragraph()
    set_para_format(p, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT,
                    first_indent=Inches(-0.25))
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(f"{num}. {text}")
    apply_font(run)
    return p


def add_finding_insight(doc, finding_text, insight_text):
    """Italic Finding / Insight blocks matching sample report format."""
    p1 = doc.add_paragraph()
    set_para_format(p1, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_indent=Inches(0.5))
    r1a = p1.add_run("Finding")
    apply_font(r1a, italic=True, bold=False)
    r1b = p1.add_run(": " + finding_text)
    apply_font(r1b, italic=True)

    p2 = doc.add_paragraph()
    set_para_format(p2, space_after=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_indent=Inches(0.5))
    r2a = p2.add_run("Insight")
    apply_font(r2a, italic=True, bold=False)
    r2b = p2.add_run(": " + insight_text)
    apply_font(r2b, italic=True)


def add_figure(doc, fig_path, caption=None, width=5.5):
    doc.add_picture(str(fig_path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(cp, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = cp.add_run(caption)
        apply_font(run, size=10)


# ═════════════════════════════════════════════════════════════════════════════
# Build Document
# ═════════════════════════════════════════════════════════════════════════════
doc = Document()

for style in doc.styles:
    if style.name == "Normal":
        style.font.name = "Arial"
        style.font.size = Pt(12)

section = doc.sections[0]
set_margins(section)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
def cover_line(doc, text, bold=False, size=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, space_after=space_after, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    apply_font(run, bold=bold, size=size)
    return p

cover_line(doc, " ", space_after=24)

cover_line(doc, "Transaction Patterns and Fraud Risk Indicators in Digital Banking:",
           bold=True, size=14, space_after=4)
cover_line(doc, "A 2024 Data-Driven Analysis", bold=True, size=14, space_after=48)

cover_line(doc, "Submitted by:", bold=True, space_after=6)
cover_line(doc, "[Member 1]", space_after=2)
cover_line(doc, "[Member 2]", space_after=2)
cover_line(doc, "[Member 3]", space_after=24)

cover_line(doc, "Instructor: [Instructor Name]", space_after=4)
cover_line(doc, "Date: April 2026", space_after=0)

doc.add_page_break()

# ── I. INTRODUCTION ───────────────────────────────────────────────────────────
add_heading(doc, "I. Introduction")
add_body(doc,
    "Digital banking has changed how people and organizations handle money. Mobile payments, "
    "contactless cards, and online transfers now account for the majority of transactions in most "
    "markets. The Association of Certified Fraud Examiners estimates that financial institutions "
    "lose billions of dollars annually to payment fraud, and digital channels are where an "
    "increasing share of that fraud occurs."
)
add_body(doc,
    "Banks and payment processors have relied on rule-based detection systems for decades, but "
    "static rules struggle to keep pace with how fraud tactics evolve. Identifying what actually "
    "separates fraudulent transactions from legitimate ones requires examining the transaction data "
    "directly, without assuming which variables matter in advance."
)

add_heading(doc, "Importance of the Study")
add_body(doc,
    "Fraud detection in digital banking has direct financial and societal consequences. Financial "
    "institutions that cannot distinguish fraudulent from legitimate transactions incur direct "
    "monetary losses, damage customer trust, and face regulatory penalties. For customers, "
    "undetected fraud causes financial harm and disrupts access to banking services. As digital "
    "transaction volumes grow and fraud tactics evolve, data-driven approaches that can identify "
    "patterns across millions of transactions are increasingly essential to maintaining the "
    "integrity of digital financial infrastructure."
)

add_heading(doc, "Research Question")
add_body(doc,
    "What transaction characteristics and customer behaviors most significantly distinguish "
    "fraudulent transactions from legitimate ones in digital banking platforms?"
)

add_heading(doc, "Objectives")
add_numbered(doc, 1, "Identify the merchant categories, transaction methods, and time windows with the highest fraud rates.")
add_numbered(doc, 2, "Examine whether customer demographic variables (age, income) are associated with fraud victimization.")
add_numbered(doc, 3, "Determine which card attributes (card type, chip technology) are most protective against fraud.")

# ── II. DATASET DESCRIPTION ───────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "II. Dataset Description")
add_body(doc,
    "The primary dataset used in this study is the \"Financial Transactions Dataset: Analytics,\" "
    "published on Kaggle by computingvictor [1]. It simulates a realistic digital banking "
    "environment and is distributed across five interrelated files covering transactions, customer "
    "demographics, card attributes, fraud labels, and merchant category codes."
)

add_label_line(doc, "Date collected: October 2024")
add_label_line(doc, "Source: Kaggle (https://www.kaggle.com), computingvictor")
add_label_line(doc, "Labelled transactions: 8,914,963 records with fraud/legitimate classification")
add_label_line(doc, "Fraud transactions: 13,332 (0.15% overall fraud rate)")
add_label_line(doc, "Date range: Multi-year, from 2010 onward")

doc.add_paragraph()
add_body(doc, "The dataset is distributed across the following five files:")

table = doc.add_table(rows=6, cols=3)
table.style = "Table Grid"
headers = ["File", "Records", "Key Variables"]
rows_data = [
    ("transactions_data.csv", "13,305,915",
     "id, date, client_id, card_id, amount, use_chip, merchant_id, mcc, errors"),
    ("users_data.csv", "2,000",
     "id, current_age, gender, yearly_income, credit_score, num_credit_cards"),
    ("cards_data.csv", "6,146",
     "id, client_id, card_brand, card_type, has_chip, credit_limit"),
    ("train_fraud_labels.json", "8,914,963",
     "Transaction ID to fraud label (Yes/No)"),
    ("mcc_codes.json", "150 codes",
     "MCC integer code to merchant category description"),
]
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for run in hdr_cells[i].paragraphs[0].runs:
        apply_font(run, bold=True, size=11)
for row_idx, row_data in enumerate(rows_data, start=1):
    cells = table.rows[row_idx].cells
    for col_idx, val in enumerate(row_data):
        cells[col_idx].text = val
        for run in cells[col_idx].paragraphs[0].runs:
            apply_font(run, size=10)

doc.add_paragraph()
add_heading(doc, "Limitations")
add_numbered(doc, 1, "The data is synthetically generated and may not capture the full complexity of real-world fraud patterns.")
add_numbered(doc, 2, "Approximately 33% of transactions had no fraud label and were excluded from the labelled analysis.")
add_numbered(doc, 3, "Merchant city and state variables were excluded due to high cardinality.")
add_numbered(doc, 4, "The dataset contains no behavioral signals such as transaction velocity or device fingerprinting.")

# ── III. METHODOLOGY ─────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "III. Methodology")
add_body(doc,
    "The analysis was conducted using Python 3.12 as the primary programming language throughout. "
    "The Pandas library was used for data cleaning, transformation, and statistical summarization. "
    "Matplotlib and Seaborn were employed to produce the six visualizations included in this report. "
    "The report itself was assembled programmatically using python-docx to ensure compliance with "
    "the prescribed A4, Arial 12, 1.5-line-spacing format specification."
)
add_body(doc,
    "The analytical process began with data ingestion and memory optimization. Because the primary "
    "transactions file is approximately 1.2 GB, column data types were explicitly set during loading "
    "— integer fields as int32, numeric fields as float32, and categorical fields as the pandas "
    "category type — reducing memory use by approximately 60%. Fraud labels were loaded from a "
    "separate JSON file and joined to the transactions table via an inner join on transaction ID, "
    "producing a labelled dataset of 8,914,963 records."
)
add_body(doc,
    "Data cleaning and feature engineering followed. Dollar signs and commas were stripped from "
    "monetary fields (amount, yearly_income, credit_limit) and converted to numeric values. "
    "Transaction dates were parsed and decomposed into hour-of-day, day-of-week, and month "
    "attributes to enable temporal analysis. Age and income were binned into categorical groups "
    "for demographic segmentation. MCC integer codes were replaced with readable merchant category "
    "descriptions using the dataset's dedicated lookup file."
)
add_body(doc,
    "Exploratory data analysis was then performed across all primary segmentation variables: "
    "merchant categories, transaction methods, card types, customer demographics, and temporal "
    "dimensions. Fraud rates were calculated as the proportion of fraudulent transactions within "
    "each segment. Summary statistics — including medians and means — were computed for transaction "
    "amounts, stratified by fraud status, to characterize the distributional differences between "
    "legitimate and fraudulent activity."
)
add_body(doc,
    "Finally, six high-quality visualizations were generated at 300 DPI to communicate the key "
    "analytical findings: a merchant category fraud rate chart, a log-scale amount distribution "
    "histogram, a fraud rate heatmap by hour and day of week, demographic bar charts by age group "
    "and income bracket, card type and transaction method charts, and a Pearson correlation matrix "
    "across all numerical features. Correlation analysis was conducted to examine linear "
    "relationships between numerical features and the fraud label, confirming the need for "
    "multivariate approaches in any subsequent fraud detection model."
)

# ── IV. DATA ANALYSIS & VISUAL FINDINGS ──────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "IV. Data Analysis and Visual Findings")

# --- Fig 1
add_heading(doc, "A. Fraud Rate by Merchant Category (Top 15 Highest-Risk)")
add_figure(doc, FIGS / "fig1_fraud_by_mcc.png",
           "Figure 1. Top 15 merchant categories ranked by fraud rate.")
add_finding_insight(doc,
    "Computer and peripheral equipment retailers recorded the highest fraud rate at 10.83%, "
    "followed by electronics stores (8.57%) and precious stones and metals dealers (6.87%). "
    "These three categories are more than 45 times the dataset average of 0.15%. Routine "
    "categories such as grocery stores and service stations remained below 0.05%.",
    "High-value, easily resalable goods attract disproportionate fraud activity. Merchants in "
    "computer equipment and electronics categories should face lower transaction limits and "
    "mandatory step-up authentication for purchases above $500. Acquiring banks can reduce a "
    "substantial share of total fraud losses by applying stricter controls specifically to "
    "these MCC codes rather than across all merchant categories uniformly."
)

# --- Fig 2
add_heading(doc, "B. Transaction Amount Distribution — Fraud vs. Legitimate")
add_figure(doc, FIGS / "fig2_amount_distribution.png",
           "Figure 2. Log-scale histogram of transaction amounts by fraud status.")
add_finding_insight(doc,
    "Legitimate transactions are concentrated in the $10 to $100 range with a median of $28.95. "
    "Fraudulent transactions have a distinctly higher median of $69.97 and a longer right tail "
    "extending into the hundreds and low thousands of dollars. The fraud distribution is shifted "
    "noticeably to the right relative to the legitimate distribution, with elevated density in "
    "the $100 to $1,000 range that is not present in the legitimate curve.",
    "Fraudulent transactions skew toward higher amounts, which is consistent with an intent to "
    "maximize return before detection occurs. Transaction amount alone is a weak standalone "
    "predictor — its value is greatest when used in combination with merchant category and "
    "transaction method, where high-amount transactions at high-risk MCC codes provide a "
    "compound signal worth acting on."
)

# --- Fig 3
add_heading(doc, "C. Fraud Rate Heatmap — Hour of Day by Day of Week")
add_figure(doc, FIGS / "fig3_fraud_heatmap_time.png",
           "Figure 3. Heatmap of fraud rate (%) by hour of day and day of week.")
add_finding_insight(doc,
    "Fraud rates are elevated from midnight to 5:00 AM across all days of the week, with the "
    "highest concentrations occurring on weekends between 1:00 AM and 4:00 AM. Weekday business "
    "hours from 9:00 AM to 5:00 PM show the lowest fraud rates, reflecting periods of active "
    "legitimate cardholder engagement.",
    "Time of day is a low-cost and reliable fraud signal. A transaction at 3:00 AM from a "
    "computer equipment retailer warrants scrutiny even if the amount appears ordinary. Financial "
    "institutions can trigger soft authentication challenges — push notifications or one-time "
    "passcodes — during off-peak hours rather than issuing hard declines, preserving access "
    "for legitimate late-night users while increasing detection of fraudulent activity."
)

# --- Fig 4
add_heading(doc, "D. Fraud Rate by Customer Demographics")
add_figure(doc, FIGS / "fig4_fraud_by_demographic.png",
           "Figure 4. Fraud rate by age group (left) and annual income bracket (right).")
add_finding_insight(doc,
    "Customers under 25 years of age exhibited the highest fraud rate among all age groups, while "
    "the 55 to 64 cohort showed the lowest susceptibility. Across income brackets, customers "
    "earning under $30,000 annually experienced the highest fraud rates, with rates declining "
    "progressively as income increased.",
    "Younger and lower-income customers are disproportionately affected by fraud, likely due to "
    "less security-conscious digital behavior and a lower likelihood of noticing unfamiliar charges "
    "quickly. Financial institutions should prioritize proactive fraud education, in-app alerts "
    "for unusual activity, and simplified fraud reporting channels specifically for these "
    "demographic segments."
)

# --- Fig 5
add_heading(doc, "E. Fraud Rate by Card Type and Transaction Method")
add_figure(doc, FIGS / "fig5_fraud_by_card_type.png",
           "Figure 5. Fraud rate by transaction method (left), card type (center), and chip availability (right).")
add_finding_insight(doc,
    "Online transactions recorded a fraud rate of 0.84%, which is 28 times higher than swipe "
    "transactions (0.03%) and 8 times higher than chip-authenticated transactions (0.10%). "
    "Prepaid debit cards exhibited a higher fraud rate (0.22%) compared to standard debit "
    "(0.13%) and credit cards (0.16%). Cards without chip technology showed markedly higher "
    "fraud rates than chip-enabled cards, confirming the protective value of EMV technology.",
    "The card-not-present environment of online transactions is the dominant fraud vector in "
    "this dataset. Investment in 3D Secure 2.0 protocols, behavioral biometrics for online "
    "sessions, and enhanced KYC verification for prepaid card issuance are the highest-impact "
    "interventions available. Accelerating the phase-out of non-chip cards, with prepaid "
    "customers prioritized, would directly reduce counterfeit card fraud at physical terminals."
)

# --- Fig 6
add_heading(doc, "F. Correlation Matrix of Numerical Features")
add_figure(doc, FIGS / "fig6_correlation_heatmap.png",
           "Figure 6. Pearson correlation matrix of numerical features including the is_fraud label.")
add_finding_insight(doc,
    "No single numerical feature shows a strong linear correlation with the fraud label. "
    "Hour-of-day has the highest positive correlation at r = +0.04, consistent with the temporal "
    "patterns identified in Figure 3. Transaction amount shows a modest positive correlation, "
    "while credit score shows a slight negative correlation with fraud incidence. Feature "
    "inter-correlations are generally low, indicating limited multicollinearity.",
    "The absence of strong individual correlations confirms that fraud cannot be detected "
    "reliably through any single threshold rule. This result supports the use of multivariate "
    "modeling — particularly tree-based ensemble methods such as Gradient Boosting or Random "
    "Forest, which can capture non-linear interactions between features. Augmenting the feature "
    "set with velocity metrics and merchant-level deviation scores would significantly strengthen "
    "any downstream classification model."
)

# ── V. KEY INSIGHTS ──────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "V. Key Insights")
add_body(doc,
    "The analysis of 8,914,963 labelled digital banking transactions reveals that fraud is not "
    "randomly distributed across the dataset. It is consistently concentrated along specific, "
    "measurable dimensions — merchant category, transaction channel, time of day, and card "
    "technology — each of which represents an actionable lever for fraud reduction."
)
add_body(doc,
    "Online transactions are the dominant fraud channel, accounting for a fraud rate of 0.84% "
    "despite representing only 11.7% of total transaction volume. This rate is 28 times higher "
    "than swipe transactions and 8 times higher than chip-authenticated transactions. The "
    "card-not-present environment is the single largest source of fraud exposure in the dataset "
    "and represents the clearest opportunity for targeted intervention."
)
add_body(doc,
    "Merchant category is the strongest segmentation variable identified in the analysis. "
    "Computer and peripheral equipment retailers reach fraud rates above 10%, more than 60 times "
    "the dataset average of 0.15%. Electronics stores and precious metals dealers follow closely. "
    "Fraud is not evenly distributed across merchant types — it is heavily concentrated in "
    "categories that sell high-value, easily liquidated goods, a pattern that has direct "
    "implications for MCC-level risk policy."
)
add_body(doc,
    "Temporal patterns provide a low-cost and reliable fraud signal. Fraud rates are "
    "consistently elevated between midnight and 5:00 AM across all days, with the highest "
    "concentrations on weekend overnight hours. Weekday business hours show the lowest fraud "
    "rates. Because these features are derived from existing transaction timestamps, they can "
    "be incorporated into real-time scoring models at no additional data collection cost."
)
add_body(doc,
    "Card technology has a demonstrable protective effect. Non-chip cards show significantly "
    "higher fraud rates than EMV chip-enabled cards, consistent with the well-documented "
    "effectiveness of chip technology in suppressing counterfeit card fraud at physical "
    "terminals. The remaining gap is concentrated in online channels, where chip technology "
    "offers no protection, reinforcing the priority of investment in digital authentication."
)
add_body(doc,
    "Transaction amount provides a useful supporting signal. The median fraudulent transaction "
    "of $69.97 is more than double the median legitimate transaction of $28.95, and the fraud "
    "distribution has a longer right tail extending into the hundreds and low thousands of "
    "dollars. Amount is a weak standalone predictor but becomes a meaningful input when "
    "combined with merchant category and transaction method in a composite risk score."
)

# ── VI. RECOMMENDATIONS ──────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "VI. Recommendations")
add_body(doc,
    "Based on the findings of this analysis, five strategic actions are recommended to address "
    "the highest-risk fraud patterns identified in the dataset. Each recommendation is grounded "
    "directly in the quantitative results and targets a specific, measurable source of fraud exposure."
)
add_body(doc,
    "First, financial institutions should implement MCC-specific transaction controls. Requiring "
    "step-up authentication for transactions above $500 at the highest-risk merchant category "
    "codes — computer equipment retailers (MCC 5045), electronics stores (MCC 5065), and precious "
    "metals dealers (MCC 5094) — would directly target the categories where fraud rates exceed "
    "10%, without creating friction across the vast majority of low-risk merchant transactions. "
    "Dynamic limit structures calibrated to each category's baseline fraud rate would be more "
    "effective than a uniform policy applied across all merchants."
)
add_body(doc,
    "Second, online channel security requires dedicated and sustained investment. Online "
    "transactions are 28 times more likely to be fraudulent than swipe transactions, and their "
    "share of total transaction volume is growing. All card-not-present transactions should be "
    "routed through 3D Secure 2.0. Layering behavioral biometrics — typing cadence, device "
    "fingerprinting, and session anomaly detection — adds meaningful protection without "
    "significantly increasing friction for legitimate users. Prepaid card programs, which showed "
    "elevated fraud rates relative to standard debit and credit, should face enhanced identity "
    "verification requirements at the point of issuance."
)
add_body(doc,
    "Third, time-of-day risk scoring should be integrated into real-time fraud detection models. "
    "Hour-of-day and day-of-week are reliable signals that can be derived from existing "
    "transaction timestamps at no additional collection cost. Transactions between midnight and "
    "5:00 AM, particularly on weekends, should receive a higher composite risk score that "
    "triggers a soft authentication challenge — a push notification or one-time passcode — "
    "rather than a hard decline. This preserves access for legitimate late-night users while "
    "increasing scrutiny during the windows of greatest fraud concentration."
)
add_body(doc,
    "Fourth, the phase-out of non-chip cards should be accelerated. Cards without EMV chip "
    "technology show substantially higher fraud rates than chip-enabled alternatives, and "
    "counterfeit card fraud at physical terminals remains a measurable contributor to overall "
    "losses. Financial institutions should proactively issue chip card replacements to remaining "
    "magnetic-stripe-only cardholders, with prepaid debit customers prioritized given their "
    "elevated fraud exposure as identified in this analysis."
)
add_body(doc,
    "Fifth, fraud awareness and monitoring programs should be targeted toward younger and "
    "lower-income customer segments. Customers under 25 and those earning under $30,000 annually "
    "showed above-average fraud victimization rates. Proactive in-app alerts for unusual "
    "transaction patterns, personalized notifications when purchases occur in high-risk contexts, "
    "and a streamlined fraud reporting process would reduce both the volume of undetected fraud "
    "and the time required to identify and resolve it."
)

# ── VII. CONCLUSION ───────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "VII. Conclusion")
add_body(doc,
    "This study examined 8,914,963 labelled digital banking transactions to identify the "
    "transaction characteristics and customer behaviors most significantly associated with fraud. "
    "The findings confirm that fraud in digital banking is not randomly distributed. It clusters "
    "in specific merchant categories, transaction channels, time windows, and card types, each "
    "of which can be measured and targeted through data-driven controls."
)
add_body(doc,
    "The online channel emerges as the most critical vulnerability, with fraud rates 28 times "
    "higher than physical swipe transactions. Merchant category is the strongest single "
    "discriminator, with computer equipment retailers reaching fraud rates above 10% against "
    "a dataset average of 0.15%. EMV chip technology demonstrably suppresses fraud at physical "
    "terminals, and temporal patterns — particularly the elevated fraud during late-night weekend "
    "hours — provide a reliable, zero-cost input for real-time risk scoring."
)
add_body(doc,
    "These findings confirm the research question and establish that no single variable is "
    "sufficient for fraud detection in isolation. The correlation matrix in Figure 6 shows "
    "that individual features carry weak linear signals, which means effective detection "
    "requires combining merchant category, transaction method, time of day, and card type "
    "into a multivariate risk score. Financial institutions that operationalize these insights "
    "through dynamic MCC-level controls, enhanced online channel authentication, and temporal "
    "risk scoring stand to reduce fraud losses substantially while minimizing unnecessary "
    "friction for legitimate customers."
)
add_body(doc,
    "Future work should incorporate additional behavioral features such as transaction velocity "
    "and geographic distance from prior activity, explore machine learning classification models "
    "including Gradient Boosting and Isolation Forest, and validate these findings against live "
    "production data from an actual financial institution rather than a simulated dataset."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "References")
add_body(doc,
    "[1] computingvictor, \"Financial Transactions Dataset: Analytics,\" Kaggle, Oct. 2024. "
    "[Online]. Available: https://www.kaggle.com/datasets/computingvictor/transactions-fraud-datasets. "
    "[Accessed: Apr. 12, 2026]."
)
add_body(doc,
    "[2] T. Hastie, R. Tibshirani, and J. Friedman, The Elements of Statistical Learning: "
    "Data Mining, Inference, and Prediction, 2nd ed. New York, NY, USA: Springer, 2009."
)
add_body(doc,
    "[3] A. Dal Pozzolo, O. Caelen, R. A. Johnson, and G. Bontempi, \"Calibrating probability "
    "with undersampling for unbalanced classification,\" in Proc. IEEE Symp. Comput. Intell. "
    "Data Mining (CIDM), Orlando, FL, USA, Dec. 2015, pp. 159-166."
)

# ── Save & Validate ───────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved -> {OUTPUT}")

print("\n=== VALIDATION ===")
doc2 = Document(OUTPUT)
headings = [p.text for p in doc2.paragraphs if p.runs and p.runs[0].bold and
            any(kw in p.text for kw in ["Introduction", "Dataset", "Methodology", "Analysis",
                                         "Insights", "Recommendation", "Conclusion", "Reference",
                                         "Importance", "Research Question", "Objectives", "Limitations"])]
print(f"Section/sub-section headings found: {len(headings)}")
for h in headings:
    print(f"  * {h[:80]}")
print(f"Embedded images: {len(doc2.inline_shapes)}")
s = doc2.sections[0]
print(f"Page size: {s.page_width.cm:.1f} cm x {s.page_height.cm:.1f} cm (A4 = 21.0 x 29.7)")
print(f"Margins - Top:{s.top_margin.inches:.2f}\" Right:{s.right_margin.inches:.2f}\" "
      f"Bottom:{s.bottom_margin.inches:.2f}\" Left:{s.left_margin.inches:.2f}\"")
total_words = sum(len(p.text.split()) for p in doc2.paragraphs)
print(f"Estimated word count: {total_words:,}")
